import pandas as pd
from docx import Document
import sys
import os

def load_placeholder_data(excel_path):
    df = pd.read_excel(excel_path)
    if df.shape[1] < 2:
        raise ValueError("Excel must have at least two columns: Placeholder and Value")
    return dict(zip(df.iloc[:, 0], df.iloc[:, 1].astype(str)))

def replace_placeholders(doc, data_dict):
    for para in doc.paragraphs:
        for key, value in data_dict.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data_dict.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

def generate_document(template_path, excel_path, output_path):
    data = load_placeholder_data(excel_path)
    doc = Document(template_path)
    replace_placeholders(doc, data)
    doc.save(output_path)
    print(f"✅ Document generated: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python generate_doc_from_excel.py <template.docx> <data.xlsx> <output.docx>")
        sys.exit(1)

    template_file = sys.argv[1]
    excel_file = sys.argv[2]
    output_file = sys.argv[3]

    if not os.path.exists(template_file):
        print(f"❌ Template file not found: {template_file}")
        sys.exit(1)

    if not os.path.exists(excel_file):
        print(f"❌ Excel file not found: {excel_file}")
        sys.exit(1)

    generate_document(template_file, excel_file, output_file)

