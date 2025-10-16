from docx import Document
import io

def generate_document(data_dict, template_file):
    doc = Document(template_file)

    for para in doc.paragraphs:
        inline_replace(para, data_dict)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                inline_replace(cell, data_dict)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_preview(data_dict, template_file):
    doc = Document(template_file)
    preview = []

    for para in doc.paragraphs:
        text = para.text
        for key, value in data_dict.items():
            text = text.replace(f"{{{{{key}}}}}", value)
        if text.strip():
            preview.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for key, value in data_dict.items():
                    text = text.replace(f"{{{{{key}}}}}", value)
                if text.strip():
                    preview.append(text)

    return "\n".join(preview)

def inline_replace(container, data_dict):
    for key, value in data_dict.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in container.text:
            container.text = container.text.replace(placeholder, value)

